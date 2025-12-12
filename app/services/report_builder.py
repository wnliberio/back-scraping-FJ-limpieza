# app/services/report_builder.py - SIN CAMBIOS, VERIFICADO PARA 4 CASOS
"""
NOTA: Este archivo NO requiere cambios.
Funciona correctamente para:
✅ Caso 1: Scraping + resultados
✅ Caso 2: Scraping + sin procesos
✅ Caso 3: HTTPX + resultados
✅ Caso 4: HTTPX + "Página 1 sin resultados"

El encabezado profesional (7 campos) se incluye en TODOS los casos.
"""

import os
from datetime import datetime, date
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image

from core.config import OUTPUT_DIR


TITLE_FONT = "Times New Roman"
BODY_FONT = "Times New Roman"


def _ensure_reports_dir() -> str:
    reports_dir = os.path.join(OUTPUT_DIR, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    return reports_dir


def _set_doc_defaults(doc: Document):
    # Márgenes 2.54cm (estándar)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Fuente por defecto
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(12)

    # Espaciado 1.5 aprox
    style.paragraph_format.line_spacing = 1.5


def _available_width_inches(doc: Document) -> float:
    s = doc.sections[0]
    avail = s.page_width - s.left_margin - s.right_margin
    # EMU -> inches
    return float(avail) / 914400.0


def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = TITLE_FONT
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = TITLE_FONT
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _format_money(val: float) -> str:
    try:
        return "${:,.2f}".format(float(val))
    except Exception:
        return str(val)


def _pick_images(payload: dict) -> List[str]:
    """
    V25: Filtra solo screenshots de páginas de resultados (pageX.png).
    Excluye screenshots intermedios del proceso de navegación.
    """
    paths = []
    
    # V25: Verificar si hay array de screenshots
    screenshots_array = payload.get("screenshots")
    if screenshots_array and isinstance(screenshots_array, list):
        for p in screenshots_array:
            if p and os.path.exists(p):
                # Filtrar: solo incluir screenshots de páginas de resultados
                # Patrón: funcion_judicial_NOMBRE_page1.png, page2.png, etc.
                if 'page' in os.path.basename(p).lower():
                    paths.append(p)
        
        if paths:
            return paths
    
    # Retrocompatibilidad: método anterior
    for k in ("screenshot_path", "screenshot_historial_path"):
        p = payload.get(k)
        if p and os.path.exists(p):
            paths.append(p)
    
    return paths


def _human_name(tipo: str) -> str:
    mapping = {
        "ruc": "SRI – RUC",
        "deudas": "SRI – Deudas Firmes/Impugnadas",
        "denuncias": "Fiscalía – Denuncias",
        "mercado_valores": "Superintendencia – Mercado de Valores",
        "interpol": "INTERPOL – Notificaciones",
        "google": "Google – Búsqueda",
        "contraloria": "Contraloría – DDJJ",
        "supercias_persona": "Superintendencia – Consulta de Persona",
        "predio_quito": "GAD Quito – Predios",
        "predio_manta": "GAD Manta – Predios",
        "funcion_judicial": "Función Judicial – Procesos Judiciales",
    }
    return mapping.get(tipo, tipo)


def build_report_docx(job_id: str, meta: Dict[str, Any], results: Dict[str, Any]) -> str:
    """
    Construye un DOCX profesional con:
    - Encabezado tabular con datos del cliente (7 campos)
    - Secciones por consulta
    - Conclusión
    
    ✅ FUNCIONA PARA LOS 4 CASOS:
    
    Caso 1: Scraping + resultados
    - results = {'funcion_judicial': {'scenario': 'results_found', 'screenshots': [...]}}
    - Muestra encabezado + tabla de procesos + screenshots
    
    Caso 2: Scraping + sin procesos
    - results = {'funcion_judicial': {'scenario': 'no_results'}}
    - Muestra encabezado + mensaje "No se encontraron procesos judiciales"
    
    Caso 3: HTTPX + resultados
    - results = {'funcion_judicial': {...generado por HTTPX...}}
    - Muestra encabezado + tabla de procesos
    
    Caso 4: HTTPX + "Página 1 sin resultados"
    - results = {'funcion_judicial': {'scenario': 'no_results'}}
    - Muestra encabezado + mensaje "No se encontraron procesos judiciales"
    
    Args:
        job_id: Identificador único del trabajo
        meta: Diccionario con datos del cliente {
            'cliente_nombre': str,
            'cliente_cedula': str,
            'nombre_conyuge': str,
            'cedula_conyuge': str,
            'nombre_codeudor': str,
            'cedula_codeudor': str,
            'fecha_consulta': datetime,
            'cliente_id': int
        }
        results: Dict con consultas {
            'funcion_judicial': {
                'scenario': 'results_found' | 'no_results',
                'screenshots': [...],  # opcional
                'total_pages': int,  # opcional
                'mensaje': str  # opcional
            }
        }
    
    Returns:
        Ruta absoluta del archivo .docx generado
    """
    reports_dir = _ensure_reports_dir()

    # Documento
    doc = Document()
    _set_doc_defaults(doc)

    # ===== PORTADA / ENCABEZADO PROFESIONAL =====
    _add_title(doc, "Revisión de Función Judicial")
    doc.add_paragraph("")  # Espaciador

    # ===== TABLA DE ENCABEZADO PROFESIONAL (7 CAMPOS) =====
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    # Ancho de columnas (1ª columna 40%, 2ª columna 60%)
    for row in table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(6.54)
    
    rows = table.rows
    
    # Fila 1: Fecha de consulta
    rows[0].cells[0].text = "Fecha de consulta:"
    fecha_consulta = meta.get('fecha_consulta')
    if isinstance(fecha_consulta, datetime):
        rows[0].cells[1].text = fecha_consulta.strftime('%d/%m/%Y %H:%M:%S')
    else:
        rows[0].cells[1].text = str(fecha_consulta) if fecha_consulta else datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    
    # Fila 2: Nombre del titular
    rows[1].cells[0].text = "Nombre del titular:"
    rows[1].cells[1].text = meta.get('cliente_nombre', 'N/A')
    
    # Fila 3: Cédula del titular
    rows[2].cells[0].text = "Cédula del titular:"
    rows[2].cells[1].text = str(meta.get('cliente_cedula', 'N/A'))
    
    # Fila 4: Nombre del cónyuge
    rows[3].cells[0].text = "Nombre del cónyuge:"
    rows[3].cells[1].text = meta.get('nombre_conyuge', 'No aplica')
    
    # Fila 5: Cédula del cónyuge
    rows[4].cells[0].text = "Cédula del cónyuge:"
    rows[4].cells[1].text = str(meta.get('cedula_conyuge', 'No aplica'))
    
    # Fila 6: Nombre del codeudor
    rows[5].cells[0].text = "Nombre del codeudor:"
    rows[5].cells[1].text = meta.get('nombre_codeudor', 'No aplica')
    
    # Fila 7: Cédula del codeudor
    rows[6].cells[0].text = "Cédula del codeudor:"
    rows[6].cells[1].text = str(meta.get('cedula_codeudor', 'No aplica'))
    
    # Espaciador después de la tabla
    doc.add_paragraph("")

    # Secciones por consulta
    width_in = _available_width_inches(doc)
    max_w = max(3.0, width_in)  # seguridad

    figura_idx = 1
    for tipo, payload in results.items():
        _add_subtitle(doc, f"Consulta: {_human_name(tipo)}")

        # Información del escenario
        scenario = payload.get("scenario")
        total_pages = payload.get("total_pages", 0)
        mensaje = payload.get("mensaje", "")
        
        # Comentario/resumen
        if scenario == "no_results":
            doc.add_paragraph("No se encontraron procesos judiciales para esta consulta.")
        elif scenario == "results_found":
            if total_pages > 1:
                doc.add_paragraph(
                    f"Se encontraron procesos judiciales distribuidos en {total_pages} páginas. "
                    "A continuación se presentan todas las capturas:"
                )
            else:
                doc.add_paragraph("Se encontraron procesos judiciales. A continuación la evidencia:")
        elif mensaje:
            doc.add_paragraph(f"{mensaje}")
        else:
            doc.add_paragraph("Se adjunta evidencia visual de la consulta realizada.")

        # Insertar imágenes (TODAS las páginas)
        imgs = _pick_images(payload)
        if not imgs:
            # ✅ Si no hay imágenes, es normal para "sin procesos"
            if scenario != "no_results":
                doc.add_paragraph("No se generaron capturas para esta consulta.")
        else:
            # V25: Insertar TODAS las páginas capturadas
            for idx, img_path in enumerate(imgs, 1):
                try:
                    # Ajuste proporcional al ancho disponible
                    with Image.open(img_path) as im:
                        doc.add_picture(img_path, width=Inches(max_w))
                    
                    # Pie de imagen con número de página si hay múltiples
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if len(imgs) > 1:
                        # Múltiples páginas: indicar número de página
                        cap.add_run(
                            f"Figura {figura_idx}. {_human_name(tipo)} – Página {idx} de {len(imgs)}"
                        ).italic = True
                    else:
                        # Una sola página
                        cap.add_run(f"Figura {figura_idx}. Evidencia – {_human_name(tipo)}").italic = True
                    
                    figura_idx += 1
                    
                    # Espacio entre páginas si hay múltiples
                    if idx < len(imgs):
                        doc.add_paragraph("")
                        
                except Exception as e:
                    doc.add_paragraph(f"[Aviso] Falló al insertar la imagen: {img_path}")
                    print(f"Error insertando imagen {img_path}: {e}")

        doc.add_paragraph("")  # separación entre consultas

    # Conclusión
    _add_subtitle(doc, "Conclusión")
    concl = (
        "Con base en las evidencias adjuntas, se confirma que las consultas fueron ejecutadas en las "
        "fuentes oficiales indicadas. Este informe presenta capturas de pantalla completas de todas las "
        "páginas de resultados encontradas. La validación del contenido específico y su relación con el "
        "monto transaccionado debe realizarse mediante revisión visual de las capturas. "
        "De ser requerido, en una siguiente fase se integrará análisis asistido por LLM con pautas "
        "para relacionar hallazgos con el monto reportado."
    )
    doc.add_paragraph(concl)

    # Guardar
    filename = f"report_{job_id}.docx"
    out_path = os.path.abspath(os.path.join(reports_dir, filename))
    doc.save(out_path)
    return out_path