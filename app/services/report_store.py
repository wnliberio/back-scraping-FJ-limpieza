# app/services/report_store.py
from __future__ import annotations

import os
import json
from datetime import datetime, date
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sqlalchemy import text  # <-- IMPORTANTE para SQLAlchemy 2.x

from core.utils.log import log
from app.db import engine  # usamos tu engine ya configurado

# Carpeta de salida (la misma base donde guardas screenshots)
OUTPUT_DIR = os.path.join(os.getcwd(), "sri_ruc_output")
REPORTS_DIR = os.path.join(OUTPUT_DIR, "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)


def _parse_fecha(fecha: Optional[str | date]) -> Optional[str]:
    """Normaliza fecha 'YYYY-MM-DD' o None."""
    if fecha is None:
        return None
    if isinstance(fecha, date):
        return fecha.isoformat()
    try:
        return str(fecha).split("T")[0]
    except Exception:
        return None


def _add_heading_center(doc: Document, text_: str, level: int = 0):
    p = doc.add_heading(text_, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_image(doc: Document, path: str):
    if not path or not os.path.exists(path):
        return
    try:
        doc.add_picture(path, width=Inches(6.5))
    except Exception as e:
        log(f"⚠️ No se pudo insertar imagen {path}: {e}")


def _short_analysis(tipo: str, payload: Dict[str, Any], monto: Optional[float]) -> str:
    scen = payload.get("scenario")
    base = f"Se revisó la fuente '{tipo}'. "
    if scen:
        base += f"Escenario detectado: {scen}. "
    if monto:
        base += f""
        #base += f"Análisis en relación al monto USD {monto:,.2f}. "
    return base.strip()


def build_docx_report(job_id: str, results: Dict[str, Any], meta: Dict[str, Any]) -> str:
    """Genera el .docx con todas las capturas de las consultas."""
    tipo_alerta = (meta or {}).get("tipo_alerta") or "Alerta"
    monto_usd = (meta or {}).get("monto_usd")
    fecha_alerta = _parse_fecha((meta or {}).get("fecha_alerta"))

    doc = Document()

    # Portada
    _add_heading_center(doc, "CONSULTA DE PROCESOS JUDICIALES ELECTRÓNICOS", level=0)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Inches(0.2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p.add_run(f"Tipo de alerta: {tipo_alerta}\n").bold = True
    #if monto_usd is not None:
    #    p.add_run(f"Monto (USD): {monto_usd:,.2f}\n")
    #if fecha_alerta:
    #    p.add_run(f"Fecha de alerta: {fecha_alerta}\n")
    #p.add_run(f"Job ID: {job_id}\n")
    #p.add_run(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_page_break()

    # Secciones por consulta
    for tipo, payload in (results or {}).items():
        doc.add_heading(f"Fuente: {tipo}", level=1)

        if isinstance(payload, dict):
            doc.add_paragraph(_short_analysis(tipo, payload, monto_usd))
        else:
            doc.add_paragraph("Resultado no estructurado.")

        if isinstance(payload, dict) and payload.get("screenshot_path"):
            _add_image(doc, payload["screenshot_path"])

        if isinstance(payload, dict) and payload.get("screenshot_historial_path"):
            doc.add_paragraph("--")
            _add_image(doc, payload["screenshot_historial_path"])

        doc.add_page_break()

    # Conclusión
    doc.add_heading("Advertencia", level=1)
    concl = doc.add_paragraph()
    concl.add_run(
        "Este documento se ha generado automáticamente. Por favor, revisarlo."
    )

    path = os.path.join(REPORTS_DIR, f"report_{job_id}.docx")
    doc.save(path)
    return path


def persist_report_row(
    *,
    job_id: str,
    report_path: str,
    meta: Dict[str, Any],
    results_snapshot: Dict[str, Any],
) -> int:
    """
    Inserta 1 fila en 'reports' (esquema actual: id, job_id, tipo_alerta, monto_usd,
    fecha_alerta, file_path, data_snapshot, created_at) y retorna el ID.
    """
    tipo_alerta = (meta or {}).get("tipo_alerta") or "Alerta"
    monto_usd = (meta or {}).get("monto_usd")
    fecha_alerta = _parse_fecha((meta or {}).get("fecha_alerta"))

    payload_json = json.dumps(
        {"job_id": job_id, "results": results_snapshot, "meta": meta},
        ensure_ascii=False,
    )

    sql = text(
        """
        INSERT INTO reports (job_id, tipo_alerta, monto_usd, fecha_alerta, file_path, data_snapshot)
        VALUES (:job_id, :tipo_alerta, :monto_usd, :fecha_alerta, :file_path, :data_snapshot)
        """
    )

    params = {
        "job_id": job_id,
        "tipo_alerta": tipo_alerta,
        "monto_usd": float(monto_usd) if monto_usd is not None else None,
        "fecha_alerta": fecha_alerta,
        "file_path": report_path,
        "data_snapshot": payload_json,
    }

    # Ejecutar INSERT y obtener ID de forma segura en SQLAlchemy 2.x
    with engine.begin() as conn:
        res = conn.execute(sql, params)
        rid = None
        # MySQL/PyMySQL suele exponer lastrowid
        try:
            rid = res.lastrowid  # type: ignore[attr-defined]
        except Exception:
            rid = None
        if not rid:
            # Fallback estándar en MySQL
            try:
                rid = conn.execute(text("SELECT LAST_INSERT_ID()")).scalar_one()
            except Exception:
                rid = None

    if not rid:
        # Último fallback por unicidad (job_id + file_path)
        with engine.connect() as conn2:
            rid = conn2.execute(
                text(
                    "SELECT id FROM reports "
                    "WHERE job_id = :job_id AND file_path = :file_path "
                    "ORDER BY id DESC LIMIT 1"
                ),
                {"job_id": job_id, "file_path": report_path},
            ).scalar_one_or_none()

    if not rid:
        raise RuntimeError("No se pudo obtener el ID del reporte insertado.")

    return int(rid)


def generate_and_persist_report(
    *, job_id: str, results: Dict[str, Any], meta: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Genera el DOCX y persiste la fila en 'reports'.
    Retorna SIEMPRE: {"report_id": int, "report_path": str}
    """
    # 1) Generar DOCX
    report_path = build_docx_report(job_id=job_id, results=results, meta=meta)
    log(f"📝 Reporte .docx generado en: {report_path}")

    # 2) Guardar en DB
    try:
        rep_id = persist_report_row(
            job_id=job_id,
            report_path=report_path,
            meta=meta,
            results_snapshot=results,
        )
        log(f"✅ Reporte persistido (id={rep_id})")
        return {"report_id": rep_id, "report_path": report_path}
    except Exception as e:
        log(f"⚠️ Error al persistir reporte: {e}")
        # Devolvemos al menos la ruta del archivo generado
        return {"report_id": -1, "report_path": report_path}
