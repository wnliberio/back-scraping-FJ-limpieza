# app/services/fj_httpx_fallback.py
"""
Fallback para consultas de Función Judicial usando API directa (HTTPX).
Se usa cuando el web scraping falla (no hay screenshots de resultados).

MEJORA V2:
- SIEMPRE retorna (ruta_reporte, resultado)
- Incluso cuando "Página 1 sin resultados" → genera reporte vacío
- Nunca retorna None (salvo error crítico)

Características:
- Consulta hasta 20 páginas
- Genera DOCX con tablas formateadas
- Convierte fechas UTC → Ecuador (UTC-5)
- Guarda en sri_ruc_output\reports\
"""

import httpx
from docx import Document
from datetime import datetime, timedelta, timezone
import os
from typing import Optional, List, Dict, Any, Tuple
import traceback
from app.services.word_utils import agregar_linea_clave_valor, agregar_titulo_principal, configurar_documento

# ===== CONFIGURACIÓN =====
API_BASE_URL = "https://api.funcionjudicial.gob.ec/EXPEL-CONSULTA-CAUSAS-SERVICE"
PAGE_SIZE = 10
MAX_PAGES = 20
REPORTS_DIR = "sri_ruc_output/reports"

# Crear directorio si no existe
os.makedirs(REPORTS_DIR, exist_ok=True)


def log(msg: str):
    """Logging con timestamp para HTTPX fallback"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[HTTPX FALLBACK {timestamp}] {msg}")


def _convertir_fecha_utc_a_ecuador(fecha_str: str) -> str:
    """
    Convierte fecha UTC a hora de Ecuador (UTC-5).
    
    Args:
        fecha_str: Fecha en formato ISO (ej: "2025-11-17T00:00:00")
        
    Returns:
        Fecha formateada como dd/mm/yyyy
    """
    try:
        if not fecha_str:
            return "N/A"
        
        # Parsear fecha
        if "T" in fecha_str:
            dt_utc = datetime.fromisoformat(fecha_str.replace("Z", "+00:00"))
        else:
            dt_utc = datetime.fromisoformat(fecha_str)
        
        # Asegurar que tenga timezone UTC
        if dt_utc.tzinfo is None:
            dt_utc = dt_utc.replace(tzinfo=timezone.utc)
        
        # Convertir a Ecuador (UTC-5)
        dt_ec = dt_utc - timedelta(hours=5)
        
        # Retornar formateado
        return dt_ec.strftime("%d/%m/%Y")
    except Exception as e:
        log(f"⚠️ Error convirtiendo fecha '{fecha_str}': {e}")
        return fecha_str[:10] if len(fecha_str) >= 10 else "N/A"


def _consultar_pagina_api(nombre_buscado: str, page: int) -> Optional[List[Dict[str, Any]]]:
    """
    Consulta una página de la API de Función Judicial.
    
    Args:
        nombre_buscado: Nombre del demandado a buscar
        page: Número de página (1-based)
        
    Returns:
        Lista de procesos encontrados, lista vacía si sin resultados, None si error
    """
    try:
        url = (
            f"{API_BASE_URL}/"
            f"api/consulta-causas/informacion/buscarCausas"
            f"?page={page}&size={PAGE_SIZE}"
        )
        
        payload = {
            "numeroCausa": "",
            "actor": {
                "cedulaActor": "",
                "nombreActor": ""
            },
            "demandado": {
                "cedulaDemandado": "",
                "nombreDemandado": nombre_buscado
            },
            "provincia": "",
            "numeroFiscalia": "",
            "recaptcha": "",
            "first": page,
            "pageSize": PAGE_SIZE
        }
        
        headers = {
            "Content-Type": "application/json",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        
        # Usar httpx con timeout
        with httpx.Client(timeout=30.0) as client:
            response = client.post(url, json=payload, headers=headers)
        
        if response.status_code != 200:
            log(f"⚠️ API retornó status {response.status_code}")
            return None
        
        # Parsear respuesta
        data = response.json()
        
        # La API puede retornar dict o list
        if isinstance(data, dict):
            resultados = data.get("data", [])
        elif isinstance(data, list):
            resultados = data
        else:
            return None
        
        # ✅ MEJORA: Retornar lista vacía en lugar de None cuando sin datos
        return resultados if resultados else []
        
    except httpx.TimeoutException:
        log(f"⚠️ Timeout consultando página {page}")
        return None
    except Exception as e:
        log(f"⚠️ Error consultando API página {page}: {e}")
        return None

# Función auxiliar para formatear nombre completo de encabezado
def _formatear_nombre_completo(apellidos: str, nombres: str) -> str:
    """Combina APELLIDOS + NOMBRES, retorna 'NO APLICA' si ambos vacíos"""
    apellidos = (apellidos or "").strip()
    nombres = (nombres or "").strip()
    nombre_completo = f"{apellidos} {nombres}".strip()
    return nombre_completo if nombre_completo else "NO APLICA"

def _valor_o_no_aplica(valor: str) -> str:
    """Retorna el valor si existe, sino 'NO APLICA'"""
    if valor and str(valor).strip() and str(valor).strip().upper() not in ['N/A', 'NA', 'NONE', '']:
        return str(valor).strip()
    return "NO APLICA"


def generar_reporte_httpx(
    nombre_cliente: str,
    job_id: str,
    meta: Dict[str, Any] = None
) -> Tuple[Optional[str], Dict[str, Any]]:
    """
    Genera reporte DOCX consultando API de Función Judicial directamente.
    Se usa como fallback cuando el web scraping falla.
    
    ✅ MEJORA: SIEMPRE retorna (ruta, resultado)
    - Incluso cuando "Página 1 sin resultados" → genera reporte vacío
    - Nunca retorna (None, ...) salvo error crítico
    
    Args:
        nombre_cliente: Nombre completo del cliente (ej: "PAMELA ALEXANDRA CASTRO DEL POZO")
        job_id: ID único del proceso (ej: "daemon_8d8c1f044264")
        meta: Diccionario con datos del cliente para el encabezado profesional
        
    Returns:
        Tupla (ruta_reporte, resultado_dict)
    """
    try:
        log(f"🌐 Iniciando consulta API para: {nombre_cliente}")
        
        # 1. Crear documento
        doc = Document()
        configurar_documento(doc)  # ✅ Aplicar estilos globales
        agregar_titulo_principal(doc, "Revisión de Función Judicial")
        doc.add_paragraph("")  # Espacio después del título
        
        # ===== ENCABEZADO PROFESIONAL (7 CAMPOS CLAVE-VALOR) =====
        if meta:
            # Fecha de consulta
            agregar_linea_clave_valor(doc, "FECHA DE CONSULTA", datetime.now().strftime("%d/%m/%Y"))
            
            # Titular
            agregar_linea_clave_valor(doc, "NOMBRE Y APELLIDO DEL TITULAR", _valor_o_no_aplica(meta.get('cliente_nombre')))
            agregar_linea_clave_valor(doc, "NUMERO DE CEDULA DEL TITULAR", _valor_o_no_aplica(meta.get('cliente_cedula')))
            
            # Cónyuge: combinar APELLIDOS + NOMBRES
            nombre_conyuge_completo = _formatear_nombre_completo(
                meta.get('apellidos_conyuge', ''),
                meta.get('nombres_conyuge', '')
            )
            agregar_linea_clave_valor(doc, "NOMBRE DEL CONYUGE", nombre_conyuge_completo)
            agregar_linea_clave_valor(doc, "CEDULA DEL CONYUGE", _valor_o_no_aplica(meta.get('cedula_conyuge')))
            
            # Codeudor: combinar APELLIDOS + NOMBRES
            nombre_codeudor_completo = _formatear_nombre_completo(
                meta.get('apellidos_codeudor', ''),
                meta.get('nombres_codeudor', '')
            )
            agregar_linea_clave_valor(doc, "NOMBRE DE CODEUDOR", nombre_codeudor_completo)
            agregar_linea_clave_valor(doc, "CEDULA DEL CODEUDOR", _valor_o_no_aplica(meta.get('cedula_codeudor')))
        else:
            # Fallback: si no hay meta, usar datos mínimos
            agregar_linea_clave_valor(doc, "FECHA DE CONSULTA", datetime.now().strftime("%d/%m/%Y"))
            agregar_linea_clave_valor(doc, "NOMBRE Y APELLIDO DEL TITULAR", nombre_cliente)
            agregar_linea_clave_valor(doc, "NUMERO DE CEDULA DEL TITULAR", "NO APLICA")
            agregar_linea_clave_valor(doc, "NOMBRE DEL CONYUGE", "NO APLICA")
            agregar_linea_clave_valor(doc, "CEDULA DEL CONYUGE", "NO APLICA")
            agregar_linea_clave_valor(doc, "NOMBRE DE CODEUDOR", "NO APLICA")
            agregar_linea_clave_valor(doc, "CEDULA DEL CODEUDOR", "NO APLICA")
        
        doc.add_paragraph("")  # Espacio después del encabezado
        
        # 2. Recorrer páginas
        contador = 1
        total_resultados = 0
        alguna_pagina_con_datos = False
        pagina_actual = 1
        
        for page in range(1, MAX_PAGES + 1):
            log(f"📄 Consultando página {page}...")
            
            resultados = _consultar_pagina_api(nombre_cliente, page)
            
            # ✅ MEJORA: Diferenciar entre "sin datos" (lista vacía) y "error" (None)
            if resultados is None:
                log(f"⚠️ Error consultando página {page}, deteniendo...")
                break  # Error de red, detener
            
            if not resultados:
                # Lista vacía = "Página sin resultados" (no error)
                log(f"📭 Página {page} sin resultados, finalizando")
                break  # Sin más páginas
            
            # Sí hay datos
            alguna_pagina_con_datos = True
            pagina_actual = page
            total_resultados += len(resultados)
            
            # 3. Agregar encabezado de página
            doc.add_heading(f"Página {page}", level=2)
            
            # 4. Crear tabla con 5 columnas
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            
            # Encabezados
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "No."
            hdr_cells[1].text = "Fecha de ingreso"
            hdr_cells[2].text = "No. proceso"
            hdr_cells[3].text = "Acción / Infracción"
            hdr_cells[4].text = "Movimientos del Proceso"
            
            # 5. Llenar tabla con datos
            for caso in resultados:
                row = table.add_row().cells
                
                # Convertir fecha
                fecha_api = caso.get("fechaIngreso", "")
                fecha_formato = _convertir_fecha_utc_a_ecuador(fecha_api) if fecha_api else "N/A"
                
                # Llenar celdas
                row[0].text = str(contador)
                row[1].text = fecha_formato
                row[2].text = str(caso.get("idJuicio", "N/A"))
                row[3].text = str(caso.get("nombreDelito", "N/A"))
                row[4].text = "Movimientos del Proceso"  # Columna fija para todos
                
                contador += 1
            
            # Espacio entre páginas
            doc.add_paragraph("")
        
        # ✅ MEJORA: Generar reporte INCLUSO sin datos
        
        # Determinar escenario
        if alguna_pagina_con_datos:
            scenario = "results_found"
            mensaje = f"Se encontraron {total_resultados} procesos judiciales en {pagina_actual} página(s)"
        else:
            scenario = "no_results"
            mensaje = "NO SE ENCONTRARON PROCESOS JUDICIALES"
            # Agregar mensaje al documento
            doc.add_paragraph(mensaje)
        
        # Guardar documento
        nombre_archivo = f"reporte_FJ_httpx_{nombre_cliente.replace(' ', '_')}_{job_id}.docx"
        ruta_completa = os.path.join(REPORTS_DIR, nombre_archivo)
        
        try:
            doc.save(ruta_completa)
            log(f"✅ Reporte DOCX generado: {ruta_completa}")
            log(f"   - Escenario: {scenario}")
            log(f"   - Total procesos: {total_resultados}")
            log(f"   - Páginas: {pagina_actual}")
        except Exception as e:
            log(f"❌ Error guardando documento: {e}")
            return None, {
                "scenario": "error",
                "total_procesos": 0,
                "total_paginas": 0,
                "mensaje": f"Error guardando documento: {str(e)}"
            }
        
        # ✅ Retornar SIEMPRE (ruta, resultado)
        resultado = {
            "scenario": scenario,
            "total_procesos": total_resultados,
            "total_paginas": pagina_actual,
            "mensaje": mensaje
        }
        
        return ruta_completa, resultado
        
    except Exception as e:
        log(f"❌ Error generando reporte HTTPX: {e}")
        traceback.print_exc()
        
        # ✅ MEJORA: Retornar (None, error_dict) en caso de error crítico
        return None, {
            "scenario": "error",
            "total_procesos": 0,
            "total_paginas": 0,
            "mensaje": f"Error crítico: {str(e)}"
        }