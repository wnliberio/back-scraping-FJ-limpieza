from docx.shared import Pt

def agregar_linea_clave_valor(doc, clave, valor):
    # Crear un párrafo nuevo
    p = doc.add_paragraph()

    # Añadir la clave en mayúsculas y negrilla
    run_clave = p.add_run(clave.upper() + ": ")
    run_clave.bold = True
    # Opcional: tamaño o fuente
    # run_clave.font.size = Pt(11)

    # Añadir el valor sin negrilla
    run_valor = p.add_run(valor)

    return p

